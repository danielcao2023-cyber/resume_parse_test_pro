from google import genai
from google.genai import types
from docx import Document

import json
import mimetypes
import openpyxl
from openpyxl.styles import Alignment
from openpyxl.utils import get_column_letter
import shutil
import time
import re
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from datetime import datetime


# =========================================
# 1) 这里改成你自己的配置
# =========================================

GEMINI_API_KEY = "AIzaSyDi5hmlT0ebgk-J3kwh5yoU3D9dC-VKyls"
INPUT_FOLDER = "/Users/caogong/Desktop/ai-test/resume_parse_test_pro/resume_input"
OUTPUT_FOLDER = "/Users/caogong/Desktop/ai-test/resume_parse_test_pro/resume_output"
MODEL_NAME = "gemini-2.5-flash"


# =========================================
# 2) 初始化
# =========================================

client = genai.Client(api_key=GEMINI_API_KEY)

input_dir = Path(INPUT_FOLDER)
output_dir = Path(OUTPUT_FOLDER)
processed_dir = input_dir / "processed"
failed_dir = input_dir / "failed"
duplicated_dir = input_dir / "duplicated"

input_dir.mkdir(parents=True, exist_ok=True)
output_dir.mkdir(parents=True, exist_ok=True)
processed_dir.mkdir(parents=True, exist_ok=True)
failed_dir.mkdir(parents=True, exist_ok=True)
duplicated_dir.mkdir(parents=True, exist_ok=True)

# 全局文件读写锁，防止多线程同时修改 Excel 导致损坏
excel_lock = threading.Lock()
# 限制并发线程数，防止触发 Gemini API 每分钟 15 次的限流 (Free Tier)
MAX_WORKERS = 3


# =========================================
# 3) 多简历提取 Prompt
# =========================================

EXTRACT_PROMPT = """
你现在要从一份中文简历文件中提取关键信息。

特别注意：
1. 一个文件里可能有 1 份简历，也可能有多份简历
2. 如果文件里有多个人的简历，必须把每个人单独拆分出来
3. 绝对不能把多个人的信息合并成一个人
4. 必须返回 JSON
5. 不要加 ```json
6. 不要加解释文字
7. 缺失字段一律返回空字符串 ""

请严格返回下面这种 JSON 结构：

{
  "resumes": [
    {
      "name": "",
      "gender": "",
      "age": "",
      "phone": "",
      "email": "",
      "education": "",
      "school": "",
      "major": "",
      "workYears": "",
      "currentCity": "",
      "expectedPosition": "",
      "expectedSalary": "",
      "lastCompany": "",
      "lastPosition": "",
      "jobStatus": "",
      "targetJob": "",
      "selfEvaluation": "",
      "languages": "",
      "remark": ""
    }
  ]
}

规则：
- 如果文件里只有 1 份简历，就在 resumes 里只放 1 个对象
- 如果文件里有多份简历，必须全部拆开
- 每个对象只对应 1 个人
- 如果某些字段无法确定，留空字符串，不要瞎编
"""


# =========================================
# 4) 公告规则：学院-专业关键词
# 可按你学校实际继续补充
# =========================================

COLLEGE_RULES = {
    "智能学院": [
        "采矿", "智能采矿", "矿业", "机械", "机械工程", "机械设计", "智能制造",
        "电气", "自动化", "控制工程", "安全工程", "应急", "机电"
    ],
    "健康学院": [
        "护理", "康复", "临床医学", "临床", "医学", "护理学", "康复治疗"
    ],
    "文旅学院": [
        "设计", "视觉传达", "数字媒体", "旅游管理", "会计", "财务管理", "艺术设计"
    ],
    "信创学院": [
        "人工智能", "大数据", "数据科学", "电子信息", "信息工程", "物联网",
        "计算机", "软件工程", "网络工程", "信息安全", "网络空间安全"
    ],
    "建工学院": [
        "智能建造", "土木", "建筑", "工程管理", "工程造价"
    ],
    "马克思主义学院": [
        "马克思", "思想政治", "中共党史", "哲学", "政治学"
    ],
    "基础教学部": [
        "数学", "应用数学", "统计", "物理", "英语", "外语"
    ],
    "体育教学部": [
        "体育", "运动训练", "体育教育"
    ],
    "心理健康教育中心": [
        "心理", "心理学", "应用心理", "积极心理"
    ]
}


# =========================================
# 5) 表头映射及查重配置
# =========================================

HEADER_MAPPING = {
    "id": "序号",
    "name": "姓名",
    "gender": "性别",
    "age": "年龄",
    "phone": "联系电话",
    "email": "邮箱",
    "education": "学历",
    "school": "毕业院校",
    "major": "专业",
    "workYears": "工作年限",
    "lastCompany": "最近公司",
    "lastPosition": "最近职位",
    "expectedSalary": "期望薪资",
    "expectedCity": "期望城市",
    "jobStatus": "求职状态",
    "targetJob": "目标职位",
    "advantages": "个人优势/自我评价",
    "languages": "语言能力",
    "inputTime": "录入时间",
    "remark": "备注",
    "source": "简历来源",
    "imageUrl": "源文件标识",
    "degreeJudge": "最高学历判断",
    "doctorFlag": "是否拥有博士学位",
    "ageCheck": "年龄评估",
    "college": "推荐学院",
    "majorKeyword": "匹配专业关键词",
    "matchLevel": "匹配等级",
    "recommendFlag": "AI推荐标识",
    "decision": "初筛结果决策",
    "policy": "政策参考",
    "aiReason": "AI判定综合理由",
    "aiScore": "判定打分"
}

# =========================================
# 6) 工具函数
# =========================================

def clean_ai_json_text(text: str) -> str:
    text = (text or "").strip()

    if text.startswith("```json"):
        text = text[7:].strip()
    elif text.startswith("```"):
        text = text[3:].strip()

    if text.endswith("```"):
        text = text[:-3].strip()

    return text


def load_docx_text(file_path: Path) -> str:
    doc = Document(str(file_path))
    parts = []

    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                parts.append(" | ".join(row_text))

    return "\n".join(parts).strip()


def load_text_file(file_path: Path) -> str:
    for enc in ("utf-8", "utf-8-sig", "gbk", "gb18030"):
        try:
            return file_path.read_text(encoding=enc).strip()
        except Exception:
            pass
    raise ValueError(f"无法读取文本文件编码：{file_path.name}")


def safe_move_path(target_dir: Path, src: Path) -> Path:
    target = target_dir / src.name
    if not target.exists():
        return target

    stem = src.stem
    suffix = src.suffix
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    return target_dir / f"{stem}_{ts}{suffix}"


def is_hidden_or_system_file(file_path: Path) -> bool:
    name = file_path.name
    if name.startswith(".") or name.startswith("~$") or name.startswith("~"):
        return True
    if name.lower() in {"thumbs.db", "desktop.ini"}:
        return True
    return False


def normalize_resumes_result(parsed):
    if isinstance(parsed, dict):
        if "resumes" in parsed and isinstance(parsed["resumes"], list):
            return [x for x in parsed["resumes"] if isinstance(x, dict)]
        return [parsed]

    if isinstance(parsed, list):
        return [x for x in parsed if isinstance(x, dict)]

    raise ValueError("模型返回格式无法识别")


def parse_age(age_text: str):
    age_text = str(age_text or "").strip()
    if not age_text:
        return None
    match = re.search(r"\d{2}", age_text)
    if match:
        try:
            return int(match.group())
        except Exception:
            return None
    return None


def is_doctor(education: str) -> bool:
    text = str(education or "").lower()
    return ("博士" in text) or ("phd" in text) or ("doctor" in text)


def match_college_and_keywords(major: str, target_job: str = "", education: str = ""):
    text = " ".join([str(major or ""), str(target_job or ""), str(education or "")]).lower()

    best_college = "待人工判断"
    matched_keywords = []

    for college, keywords in COLLEGE_RULES.items():
        current = []
        for kw in keywords:
            if kw.lower() in text:
                current.append(kw)
        if len(current) > len(matched_keywords):
            matched_keywords = current
            best_college = college

    if len(matched_keywords) >= 2:
        level = "精准匹配"
    elif len(matched_keywords) == 1:
        level = "相关匹配"
    else:
        level = "待人工判断"

    return best_college, "、".join(matched_keywords), level


def build_policy_tip(doctor_flag: str, age_check: str, recommend_college: str) -> str:
    if doctor_flag == "是":
        base = "符合博士引进政策参考范围：年龄一般不超过45周岁；博士待遇分层级执行，年薪约17万-40万，安家费/购房补贴约25万-65万，科研启动金约5万-15万，特别优秀者可一人一议。"
        if recommend_college != "待人工判断":
            base += f" 当前建议优先对接：{recommend_college}。"
        if age_check == "超出45岁":
            base += " 年龄需人工复核。"
        return base
    else:
        return "本次公告重点面向博士研究生学历学位人才，当前简历可先入库备用，由用人单位再做人工判断。"


def analyze_candidate_v2(resume: dict) -> dict:
    education = str(resume.get("education") or "")
    major = str(resume.get("major") or "")
    target_job = str(resume.get("targetJob") or resume.get("expectedPosition") or "")
    age_value = parse_age(resume.get("age") or "")

    doctor_flag = "是" if is_doctor(education) else "否"
    degree_judge = education if education else "未知"

    if age_value is None:
        age_check = "未知"
    elif age_value <= 45:
        age_check = "符合"
    else:
        age_check = "超出45岁"

    college, major_keywords, match_level = match_college_and_keywords(major, target_job, education)

    if doctor_flag == "是" and age_check == "符合" and match_level in ["精准匹配", "相关匹配"]:
        recommend_flag = "是"
        decision = "A-建议优先联系"
        ai_score = 90 if match_level == "精准匹配" else 85
    elif doctor_flag == "是" and match_level in ["精准匹配", "相关匹配"]:
        recommend_flag = "是"
        decision = "B-建议进入初筛"
        ai_score = 78
    elif doctor_flag == "否" and match_level in ["精准匹配", "相关匹配"]:
        recommend_flag = "否"
        decision = "C-入库观察"
        ai_score = 65
    else:
        recommend_flag = "否"
        decision = "D-与本次博士公告不匹配"
        ai_score = 50

    policy_tip = build_policy_tip(doctor_flag, age_check, college)

    ai_reason = f"最高学历判断：{degree_judge}；是否博士：{doctor_flag}；年龄判断：{age_check}；推荐学院：{college}；专业匹配等级：{match_level}。"

    return {
        "degreeJudge": degree_judge,
        "doctorFlag": doctor_flag,
        "ageCheck": age_check,
        "college": college,
        "majorKeyword": major_keywords,
        "matchLevel": match_level,
        "recommendFlag": recommend_flag,
        "decision": decision,
        "policy": policy_tip,
        "aiReason": ai_reason,
        "aiScore": ai_score
    }


def save_resume_to_local_excel(parsed_json: dict, source_file: Path, resume_index: int) -> dict:
    analysis = analyze_candidate_v2(parsed_json)

    data = {
        # 原有字段
        "id": "",  # 预留给序号
        "name": parsed_json.get("name") or "",
        "gender": parsed_json.get("gender") or "",
        "age": parsed_json.get("age") or "",
        "phone": parsed_json.get("phone") or "",
        "email": parsed_json.get("email") or "",
        "education": parsed_json.get("education") or "",
        "school": parsed_json.get("school") or "",
        "major": parsed_json.get("major") or "",
        "workYears": parsed_json.get("workYears") or "",
        "lastCompany": parsed_json.get("lastCompany") or "",
        "lastPosition": parsed_json.get("lastPosition") or "",
        "expectedSalary": parsed_json.get("expectedSalary") or "",
        "expectedCity": parsed_json.get("currentCity") or "",
        "jobStatus": parsed_json.get("jobStatus") or "",
        "targetJob": parsed_json.get("targetJob") or parsed_json.get("expectedPosition") or "",
        "advantages": parsed_json.get("selfEvaluation") or "",
        "languages": parsed_json.get("languages") or "",
        "inputTime": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "remark": (parsed_json.get("remark") or "").strip(),
        "source": f"批量识别-{source_file.suffix.lower().lstrip('.')}",
        "imageUrl": f"{source_file.name}#resume_{resume_index}",

        # V2 新增字段
        "degreeJudge": analysis["degreeJudge"],
        "doctorFlag": analysis["doctorFlag"],
        "ageCheck": analysis["ageCheck"],
        "college": analysis["college"],
        "majorKeyword": analysis["majorKeyword"],
        "matchLevel": analysis["matchLevel"],
        "recommendFlag": analysis["recommendFlag"],
        "decision": analysis["decision"],
        "policy": analysis["policy"],
        "aiReason": analysis["aiReason"],
        "aiScore": analysis["aiScore"]
    }

    try:
        # 使用线程锁确保同一个文件同一时间只被一条线程读写
        with excel_lock:
            excel_file = output_dir / "resumes.xlsx"
            keys = list(data.keys())
            
            if excel_file.exists():
                wb = openpyxl.load_workbook(excel_file)
                ws = wb.active
                
                headers = [cell.value for cell in ws[1]]
                phone_col_idx = headers.index(HEADER_MAPPING["phone"]) + 1 if HEADER_MAPPING["phone"] in headers else None
                email_col_idx = headers.index(HEADER_MAPPING["email"]) + 1 if HEADER_MAPPING["email"] in headers else None
                
                new_phone = str(data.get("phone") or "").strip()
                new_email = str(data.get("email") or "").strip()
                
                is_dup = False
                for row in ws.iter_rows(min_row=2, values_only=True):
                    if phone_col_idx and new_phone:
                        if str(row[phone_col_idx - 1] or "").strip() == new_phone:
                            is_dup = True
                            break
                    if email_col_idx and new_email:
                        if str(row[email_col_idx - 1] or "").strip() == new_email:
                            is_dup = True
                            break
                            
                if is_dup:
                    return {"success": False, "duplicated": True, "message": "发现相同联系方式或邮箱的历史记录"}
                    
                data["id"] = ws.max_row
            else:
                wb = openpyxl.Workbook()
                ws = wb.active
                ws.title = "简历库"
                chinese_headers = [HEADER_MAPPING.get(k, k) for k in keys]
                ws.append(chinese_headers)
                data["id"] = 1
                
            row_data = [data.get(k, "") for k in keys]
            ws.append(row_data)
            
            # 自动调整行宽与行高 (自动换行)
            for col_idx, col_cells in enumerate(ws.columns, start=1):
                max_len = 0
                for cell in col_cells:
                    cell.alignment = Alignment(wrap_text=True, vertical="center")
                    val_str = str(cell.value) if cell.value else ""
                    lines = val_str.split("\n")
                    for line in lines:
                        line_len = sum(2 if ord(char) > 255 else 1 for char in line)
                        if line_len > max_len:
                            max_len = line_len
                            
                col_letter = get_column_letter(col_idx)
                ws.column_dimensions[col_letter].width = min(max(max_len + 2, 12), 40)
                
            wb.save(excel_file)
            return {"success": True}
        
    except Exception as e:
        return {
            "success": False,
            "duplicated": False,
            "message": f"写入本地 Excel 文件失败：{str(e)}"
        }


# =========================================
# 6) 调用 Gemini
# =========================================

def generate_content_with_retry(contents, max_retries=3, wait_seconds=3):
    last_error = None

    for attempt in range(1, max_retries + 1):
        try:
            response = client.models.generate_content(
                model=MODEL_NAME,
                contents=contents,
            )
            return response
        except Exception as e:
            last_error = e
            print(f"第 {attempt} 次调用失败：{e}")
            if attempt < max_retries:
                time.sleep(wait_seconds)

    raise last_error


def analyze_image_or_pdf(file_path: Path):
    mime_type, _ = mimetypes.guess_type(str(file_path))

    if file_path.suffix.lower() == ".pdf":
        mime_type = "application/pdf"

    if not mime_type:
        raise ValueError(f"无法判断 MIME 类型：{file_path.name}")

    file_bytes = file_path.read_bytes()

    response = generate_content_with_retry(
        contents=[
            types.Part.from_bytes(
                data=file_bytes,
                mime_type=mime_type,
            ),
            EXTRACT_PROMPT,
        ],
        max_retries=3,
        wait_seconds=5
    )

    ai_text = clean_ai_json_text(response.text or "")
    parsed = json.loads(ai_text)
    return normalize_resumes_result(parsed)


def analyze_text_content(text: str):
    prompt = f"""{EXTRACT_PROMPT}

下面是简历文本内容：

{text}
"""

    response = generate_content_with_retry(
        contents=prompt,
        max_retries=3,
        wait_seconds=3
    )

    ai_text = clean_ai_json_text(response.text or "")
    parsed = json.loads(ai_text)
    return normalize_resumes_result(parsed)


def analyze_file(file_path: Path):
    suffix = file_path.suffix.lower()

    if suffix in [".jpg", ".jpeg", ".png", ".webp", ".pdf"]:
        return analyze_image_or_pdf(file_path)

    if suffix in [".txt", ".md"]:
        text = load_text_file(file_path)
        if not text:
            raise ValueError("文本文件内容为空")
        return analyze_text_content(text)

    if suffix == ".docx":
        text = load_docx_text(file_path)
        if not text:
            raise ValueError("DOCX 文档内容为空")
        return analyze_text_content(text)

    if suffix == ".doc":
        raise ValueError("暂不直接支持老式 .doc，请先另存为 .docx 或 PDF")

    raise ValueError(f"不支持的文件类型：{file_path.name}")


# =========================================
# 7) 主流程
# =========================================

def process_single_file(file_path: Path) -> dict:
    """处理单个文件的线程函数，返回统计结果"""
    print(f"\n--- [线程启动] 正在解析：{file_path.name} ---")
    result_stats = {
        "written": 0,
        "duplicated": 0,
        "failed": 0,
        "file_name": file_path.name
    }

    try:
        resumes = analyze_file(file_path)

        if not resumes:
            raise ValueError("没有识别到任何简历")

        print(f"[{file_path.name}] 共识别到 {len(resumes)} 份简历")

        for idx, resume in enumerate(resumes, start=1):
            # 将多份简历写入 Excel，内部已被 excel_lock 保护
            sheet_result = save_resume_to_local_excel(resume, file_path, idx)

            if sheet_result.get("success") is True:
                result_stats["written"] += 1
                print(f"  └ [{file_path.name}] 简历 {idx}：[{resume.get('name', '未知')}] 录入成功！")
            elif sheet_result.get("duplicated") is True:
                result_stats["duplicated"] += 1
                print(f"  └ [{file_path.name}] 简历 {idx}：[{resume.get('name', '未知')}] 表格中已存在，跳过。")
            else:
                result_stats["failed"] += 1
                print(f"  └ [{file_path.name}] 简历 {idx}：[{resume.get('name', '未知')}] 录入失败：{sheet_result.get('message', '')}")

        # 文件移动操作（因为 pathlib.rename / shutil.move 本身是原子的且目标路径加了时间戳防冲突）
        if result_stats["written"] > 0 and result_stats["failed"] == 0:
            target_path = safe_move_path(processed_dir, file_path)
            shutil.move(str(file_path), str(target_path))
            print(f"✓ [{file_path.name}] 处理完成，已移至：processed")
        elif result_stats["duplicated"] == len(resumes) and len(resumes) > 0:
            target_path = safe_move_path(duplicated_dir, file_path)
            shutil.move(str(file_path), str(target_path))
            print(f"✓ [{file_path.name}] 已重复，已移至：duplicated")
        else:
            target_path = safe_move_path(failed_dir, file_path)
            shutil.move(str(file_path), str(target_path))
            print(f"× [{file_path.name}] 存在失败简历，已移至：failed")

    except Exception as e:
        print(f"× [{file_path.name}] 解析异常：{str(e)}")
        target_path = safe_move_path(failed_dir, file_path)
        try:
            shutil.move(str(file_path), str(target_path))
        except Exception:
            pass
        result_stats["failed"] += 1

    return result_stats


def process_new_files():
    files = sorted([p for p in input_dir.iterdir() if p.is_file()])
    valid_files = [f for f in files if not is_hidden_or_system_file(f)]
    
    if not valid_files:
        return

    print("\n" + "=" * 60)
    print(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] 检测到 {len(valid_files)} 份新简历，启动 {min(len(valid_files), MAX_WORKERS)} 个线程自动处理...")

    total_written = 0
    total_duplicated = 0
    total_failed = 0

    # 使用线程池并发执行，MAX_WORKERS 避免触发 API 并发限流
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        futures = {executor.submit(process_single_file, f): f for f in valid_files}
        
        for future in as_completed(futures):
            try:
                stats = future.result()
                total_written += stats["written"]
                total_duplicated += stats["duplicated"]
                total_failed += stats.get("failed", 0)  # 这里修改以避免计算上出错
            except Exception as e:
                print(f"❌ 线程致命错误: {str(e)}")
                total_failed += 1

    print(f"\n[{datetime.now().strftime('%H:%M:%S')}] 批次处理结束。")
    print(f"统计：新增录入 {total_written} | 重复跳过 {total_duplicated} | 失败/异常数量 {total_failed}")
    print("\n【系统状态】继续监听中，请随时拖入新简历...")


def main():
    print("=" * 60)
    print("🚀 AI 招聘助手已启动 (自动监听模式) 🚀")
    print("=" * 60)
    print(f"▶ 监控文件夹：{input_dir}")
    print(f"▶ 如果有报错，请前往：{failed_dir} 查看")
    print("-" * 60)
    print("【系统状态】正在监听中...（只需将简历文件拖入上方文件夹，系统将自动识别！）")
    print("提示：在控制台按下 Ctrl + C 即可随时安全退出程序。\n")

    try:
        while True:
            process_new_files()
            time.sleep(3)
    except KeyboardInterrupt:
        print("\n\n⏹ 接收到退出信号，系统已安全停止。感谢使用！")


if __name__ == "__main__":
    main()